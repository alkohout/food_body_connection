import os
import io
import logging

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import JSONResponse
from sqlalchemy.orm import Session

from app.database import get_db
from app.api.routes.auth import get_current_user
from app.models.table_class import User, UserDocument

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/documents", tags=["documents"])

UPLOAD_DIR = os.environ.get("UPLOAD_DIR", "/tmp/user_docs")
MAX_FILE_BYTES = 20 * 1024 * 1024  # 20 MB
ALLOWED_EXTENSIONS = {".pdf", ".txt", ".docx"}


def _extract_text(file_bytes: bytes, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                return "\n".join(
                    page.extract_text() or "" for page in pdf.pages
                ).strip()
        except Exception as exc:
            logger.warning("pdfplumber failed for %s: %s", filename, exc)
            return ""

    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            texts = [p.text for p in doc.paragraphs if p.text]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text not in texts:
                            texts.append(cell.text)
            return "\n".join(texts).strip()
        except Exception as exc:
            logger.warning("python-docx failed for %s: %s", filename, exc)
            return ""

    if ext == ".txt":
        return file_bytes.decode("utf-8", errors="replace").strip()

    return ""


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    description: str = Form(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    file_bytes = await file.read()
    if len(file_bytes) > MAX_FILE_BYTES:
        raise HTTPException(status_code=400, detail="File exceeds 20 MB limit.")

    user_dir = os.path.join(UPLOAD_DIR, str(current_user.user_id))
    os.makedirs(user_dir, exist_ok=True)

    safe_name = os.path.basename(file.filename or "upload")
    file_path = os.path.join(user_dir, safe_name)

    # Append a counter if the filename already exists
    if os.path.exists(file_path):
        base, extension = os.path.splitext(safe_name)
        counter = 1
        while os.path.exists(file_path):
            file_path = os.path.join(user_dir, f"{base}_{counter}{extension}")
            counter += 1

    with open(file_path, "wb") as f:
        f.write(file_bytes)

    extracted_text = _extract_text(file_bytes, safe_name)

    doc = UserDocument(
        user_id=current_user.user_id,
        filename=os.path.basename(file_path),
        description=description,
        file_path=file_path,
        extracted_text=extracted_text or None,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    return {
        "document_id": doc.document_id,
        "filename": doc.filename,
        "description": doc.description,
        "uploaded_at": doc.uploaded_at.isoformat(),
        "text_extracted": bool(extracted_text),
    }


@router.get("")
def list_documents(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    docs = (
        db.query(UserDocument)
        .filter(UserDocument.user_id == current_user.user_id)
        .order_by(UserDocument.uploaded_at.desc())
        .all()
    )
    return [
        {
            "document_id": d.document_id,
            "filename": d.filename,
            "description": d.description,
            "uploaded_at": d.uploaded_at.isoformat(),
            "has_text": bool(d.extracted_text),
        }
        for d in docs
    ]


@router.delete("/{document_id}")
def delete_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = (
        db.query(UserDocument)
        .filter(
            UserDocument.document_id == document_id,
            UserDocument.user_id == current_user.user_id,
        )
        .first()
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found.")

    try:
        if os.path.exists(doc.file_path):
            os.remove(doc.file_path)
    except OSError as exc:
        logger.warning("Could not delete file %s: %s", doc.file_path, exc)

    db.delete(doc)
    db.commit()
    return {"deleted": document_id}
